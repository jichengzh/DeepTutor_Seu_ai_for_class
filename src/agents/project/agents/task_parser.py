# -*- coding: utf-8 -*-
"""
TaskParser — 解析参考任务书文档，提取章节结构。

支持 .docx 和 .pdf 格式。返回标准化的结构字典，供 TaskGenerator 使用。
"""

import re
from pathlib import Path
from typing import Any


SECTION_PATTERNS: dict[str, list[str]] = {
    "cover":        ["课程名称", "专业", "实习时间", "封面", "课程名"],
    "objectives":   ["课程目标", "设计目标", "学习目标", "课程背景", "背景与目标"],
    "modules":      ["模块", "阶段", "任务概述", "模块概述"],
    "details":      ["设计内容", "任务内容", "具体要求", "各模块"],
    "requirements": ["作业要求", "报告要求", "提交格式", "作业格式"],
    "deliverables": ["提交成果", "交付物", "成果要求"],
    "grading":      ["成绩考核", "评分标准", "考核方式", "成绩评定"],
    "schedule":     ["时间安排", "进度计划", "时间计划"],
    "references":   ["参考资料", "教材", "附录", "参考文献", "参考资源"],
}

# 语义推断词表（当检测到中文大纲序号但无关键词匹配时使用）
_SEMANTIC_HINTS: dict[str, list[str]] = {
    "objectives":   ["目标", "背景", "简介", "概述", "前言", "介绍"],
    "modules":      ["模块", "阶段", "任务", "内容概要", "章节概述"],
    "details":      ["设计", "实现", "详细", "具体", "方案", "内容"],
    "requirements": ["要求", "规范", "标准", "格式"],
    "deliverables": ["成果", "提交", "交付", "输出"],
    "grading":      ["考核", "评分", "评价", "分值", "分数"],
    "schedule":     ["时间", "进度", "计划", "安排", "周", "日程"],
    "references":   ["参考", "资料", "文献", "附录", "书目"],
}

# 中文大纲序号正则（标题行判定）
_CHINESE_HEADING_RE = re.compile(
    r"^(?:"
    r"[一二三四五六七八九十百]+[、.．]"           # 一、二、三、
    r"|第[一二三四五六七八九十百\d]+[章节部分]"   # 第一章、第二节
    r"|（[一二三四五六七八九十\d]+）"             # （一）（二）
    r"|\([一二三四五六七八九十\d]+\)"             # (一)(二)
    r"|[①②③④⑤⑥⑦⑧⑨⑩]"                  # 带圈数字
    r"|\d+[.、．]\s"                             # 1. 1、(后跟空格)
    r")"
)

MAX_SUMMARY_LEN = 2000


def _match_section(text: str) -> str | None:
    """Return the section key if text matches any pattern, else None."""
    for key, patterns in SECTION_PATTERNS.items():
        for p in patterns:
            if p in text:
                return key
    return None


def _is_chinese_heading(text: str) -> bool:
    """判断文本是否以中文大纲序号开头（且长度合理 <60 字）。"""
    return bool(_CHINESE_HEADING_RE.match(text.strip())) and len(text.strip()) < 60


def _infer_section_by_semantics(title: str, existing_sections: dict) -> str | None:
    """当标题含序号但无关键词匹配时，通过语义词推断章节 key。已占用的 key 跳过。"""
    for key, hints in _SEMANTIC_HINTS.items():
        if key in existing_sections:
            continue
        if any(h in title for h in hints):
            return key
    return None


class TaskParser:
    """解析参考任务书，提取文档结构。"""

    def parse_docx(self, file_path: str) -> dict[str, Any]:
        """
        解析 .docx 文件，返回结构化内容。

        Returns:
            {
                "sections": {"cover": {"title": ..., "content": ...}, ...},
                "module_count": int,
                "has_grading_table": bool,
                "raw_text": str,
            }

        Raises:
            FileNotFoundError: 文件不存在
            Exception: 文件损坏或格式不符
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        from docx import Document  # python-docx

        doc = Document(str(path))  # raises if not a valid docx

        sections: dict[str, dict[str, str]] = {}
        raw_lines: list[str] = []
        current_section: str | None = None
        current_title: str = ""
        current_content_lines: list[str] = []
        has_grading_table = False
        detected_headings: list[tuple[str, str]] = []  # (heading_text, mapped_key)

        def _flush():
            nonlocal current_section, current_content_lines, current_title
            if current_section:
                existing = sections.get(current_section)
                if existing:
                    # 同一章节多次匹配时追加内容，避免覆盖
                    existing["content"] = (existing["content"] + "\n" + "\n".join(current_content_lines)).strip()
                else:
                    sections[current_section] = {
                        "title": current_title,
                        "content": "\n".join(current_content_lines).strip(),
                    }
            current_section = None
            current_title = ""
            current_content_lines = []

        # Iterate paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                if current_content_lines:
                    current_content_lines.append("")
                continue

            raw_lines.append(text)

            # Check if this paragraph is a section heading
            is_heading = para.style.name.startswith("Heading") if para.style else False
            matched = _match_section(text)
            is_cn_heading = _is_chinese_heading(text)

            if matched and (is_heading or is_cn_heading or len(text) < 40):
                _flush()
                current_section = matched
                current_title = text
                detected_headings.append((text, matched))
            elif is_cn_heading and not matched:
                # 中文序号标题但无关键词 → 语义推断，否则归入 details
                inferred = _infer_section_by_semantics(text, sections)
                _flush()
                current_section = inferred or "details"
                current_title = text
                detected_headings.append((text, current_section))
            elif current_section:
                current_content_lines.append(text)

        _flush()

        # Check tables for grading table
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(c.text for c in row.cells)
                raw_lines.append(row_text)
                if any(kw in row_text for kw in ["优秀", "良好", "分", "评分"]):
                    has_grading_table = True
                # Try to match table content to sections
                matched = _match_section(row_text)
                if matched and matched not in sections:
                    sections[matched] = {"title": row_text[:50], "content": row_text}

        # Estimate module count from modules section content
        module_count = 0
        if "modules" in sections:
            content = sections["modules"]["content"]
            import re
            module_count = len(re.findall(r"模块[一二三四五六七八九十\d]", content))
            if module_count == 0:
                # count numbered items
                module_count = len(re.findall(r"^[1-9][.、]", content, re.MULTILINE))

        return {
            "sections": sections,
            "module_count": max(module_count, 0),
            "has_grading_table": has_grading_table,
            "raw_text": "\n".join(raw_lines),
            "detected_headings": detected_headings,
        }

    def parse_pdf(self, file_path: str) -> dict[str, Any]:
        """
        解析 .pdf 文件（使用 PyMuPDF）。

        Returns:
            同 parse_docx，scanned PDF 时返回 {"error": "scanned_pdf", "raw_text": "", "sections": {}}
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            import fitz  # PyMuPDF
        except ImportError as e:
            raise ImportError("PyMuPDF (fitz) is required for PDF parsing: pip install pymupdf") from e

        doc = fitz.open(str(path))
        raw_lines: list[str] = []

        for page in doc:
            text = page.get_text()
            if text:
                raw_lines.extend(text.splitlines())

        doc.close()

        if not any(line.strip() for line in raw_lines):
            return {"error": "scanned_pdf", "raw_text": "", "sections": {}, "module_count": 0, "has_grading_table": False}

        # Re-use text-based section matching
        sections: dict[str, dict[str, str]] = {}
        current_section: str | None = None
        current_title: str = ""
        current_content_lines: list[str] = []
        detected_headings: list[tuple[str, str]] = []

        def _flush():
            nonlocal current_section, current_content_lines, current_title
            if current_section:
                existing = sections.get(current_section)
                if existing:
                    existing["content"] = (existing["content"] + "\n" + "\n".join(current_content_lines)).strip()
                else:
                    sections[current_section] = {
                        "title": current_title,
                        "content": "\n".join(current_content_lines).strip(),
                    }
            current_section = None
            current_title = ""
            current_content_lines = []

        for line in raw_lines:
            text = line.strip()
            if not text:
                continue
            matched = _match_section(text)
            is_cn_heading = _is_chinese_heading(text)

            if matched and len(text) < 60:
                _flush()
                current_section = matched
                current_title = text
                detected_headings.append((text, matched))
            elif is_cn_heading and not matched:
                inferred = _infer_section_by_semantics(text, sections)
                _flush()
                current_section = inferred or "details"
                current_title = text
                detected_headings.append((text, current_section))
            elif current_section:
                current_content_lines.append(text)

        _flush()

        return {
            "sections": sections,
            "module_count": 0,
            "has_grading_table": False,
            "raw_text": "\n".join(l.strip() for l in raw_lines if l.strip()),
            "detected_headings": detected_headings,
        }

    def extract_structure_summary(self, parsed: dict[str, Any]) -> str:
        """
        生成结构摘要字符串，用于提示词中描述参考文档的章节结构。
        截断到 MAX_SUMMARY_LEN 字符。
        """
        if parsed.get("error"):
            return f"[解析错误: {parsed['error']}]"

        sections = parsed.get("sections", {})
        lines: list[str] = []
        for key, section in sections.items():
            title = section.get("title", key)
            content_preview = section.get("content", "")[:200]
            lines.append(f"### {title}\n{content_preview}")

        lines.append(f"\n[共 {len(sections)} 个章节，模块数: {parsed.get('module_count', 0)}]")
        summary = "\n\n".join(lines)

        if len(summary) > MAX_SUMMARY_LEN:
            summary = summary[:MAX_SUMMARY_LEN] + "\n...[截断]"

        return summary


__all__ = ["TaskParser", "SECTION_PATTERNS"]
