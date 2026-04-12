# -*- coding: utf-8 -*-
"""
TaskGenerator — 逐章节流式生成新任务书 / 课程大纲。

流程：
1. 可选 RAG 检索相关知识
2. 可选 Web 搜索获取最新资料
3. 按固定顺序逐章节流式生成
4. 导出 Markdown + docx（大纲模式额外导出 PDF）
"""

import asyncio
import time
from collections import defaultdict
from pathlib import Path
from typing import Any, Callable

import yaml

from src.logging.logger import get_logger
from src.services.llm import factory as llm_factory

logger = get_logger("TaskGenerator")

# ── 任务书章节顺序 ──────────────────��───────────────────────────────
SECTION_ORDER_ZH = [
    ("cover",        "封面信息"),
    ("objectives",   "课程背景与目标"),
    ("modules",      "模块概述"),
    ("details",      "各模块设计内容"),
    ("requirements", "作业要求"),
    ("deliverables", "提交成果"),
    ("grading",      "成绩考核"),
    ("schedule",     "时间安排"),
    ("references",   "参考资源"),
]

SECTION_ORDER_EN = [
    ("cover",        "Cover Information"),
    ("objectives",   "Background and Objectives"),
    ("modules",      "Module Overview"),
    ("details",      "Module Design Details"),
    ("requirements", "Assignment Requirements"),
    ("deliverables", "Deliverables"),
    ("grading",      "Grading Criteria"),
    ("schedule",     "Schedule"),
    ("references",   "References"),
]

# ── 课程大纲章节顺序 ───────────────────────────────────��────────────
SECTION_ORDER_SYLLABUS_ZH = [
    ("cover",              "封面信息"),
    ("objectives",         "课程简介与教学目标"),
    ("prerequisites",      "前置知识要求"),
    ("content_structure",  "课程内容及学时分配"),
    ("teaching_methods",   "教学方法与手段"),
    ("grading_scheme",     "课程考核与成绩评定"),
    ("teaching_materials", "教材与参考资料"),
    ("schedule",           "教学进度安排"),
]

SECTION_ORDER_SYLLABUS_EN = [
    ("cover",              "Cover Information"),
    ("objectives",         "Course Introduction and Teaching Objectives"),
    ("prerequisites",      "Prerequisite Knowledge Requirements"),
    ("content_structure",  "Course Content and Hours Allocation"),
    ("teaching_methods",   "Teaching Methods"),
    ("grading_scheme",     "Assessment and Grading"),
    ("teaching_materials", "Textbooks and Reference Materials"),
    ("schedule",           "Teaching Schedule"),
]

_PROMPTS_DIR = Path(__file__).resolve().parents[1] / "prompts"


def _load_prompts(language: str, mode: str = "task") -> dict[str, Any]:
    lang = "zh" if language.startswith("zh") else "en"
    filename = "syllabus_generation.yaml" if mode == "syllabus" else "task_generation.yaml"
    prompt_file = _PROMPTS_DIR / lang / filename
    with open(prompt_file, encoding="utf-8") as f:
        return yaml.safe_load(f)


class TaskGenerator:
    """
    逐章节流式生成新任务书 / 课程大纲。

    Args:
        output_dir: 生成文件存储目录
        language: 提示词语言（"zh" 或 "en"）
        mode: 生成模式，"task"（任务书）或 "syllabus"（课程大纲）
    """

    def __init__(self, output_dir: str, language: str = "zh", mode: str = "task"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.language = language
        self.mode = mode
        self._prompts = _load_prompts(language, mode)
        if mode == "syllabus":
            self._section_order = (
                SECTION_ORDER_SYLLABUS_ZH if language.startswith("zh")
                else SECTION_ORDER_SYLLABUS_EN
            )
        else:
            self._section_order = (
                SECTION_ORDER_ZH if language.startswith("zh")
                else SECTION_ORDER_EN
            )

    async def generate(
        self,
        theme: str,
        reference_structure: dict[str, Any],
        kb_name: str | None,
        web_search: bool,
        ws_callback: Callable,
    ) -> dict[str, Any]:
        """
        生成完整任务书文档。

        Returns:
            {"content": full_md, "md_path": str, "docx_path": str}
        """
        # Step 1: RAG 检索
        if kb_name:
            await ws_callback({"type": "status", "content": f"正在检索知识库 [{kb_name}]..."})
            rag_context = await self._rag_retrieve(theme, kb_name)
        else:
            await ws_callback({"type": "status", "content": "未选择知识库，无外部知识源，将仅凭 LLM 知识生成"})
            rag_context = ""

        # Step 2: Web 搜索
        web_context = ""
        if web_search:
            await ws_callback({"type": "status", "content": "正在进行网络搜索..."})
            web_context = await self._web_search(theme)

        # Step 3: 加载全局反思并注入 system prompt
        try:
            from src.agents.project.reflection_manager import get_reflection_manager
            reflection_text = get_reflection_manager().get_prompt_text()
        except Exception:
            reflection_text = ""

        # Step 4: 逐章节流式生成
        sections: dict[str, str] = {}
        base_system = self._prompts.get("system", "You are a helpful assistant.")
        system_prompt = f"{base_system}\n\n{reflection_text}" if reflection_text else base_system
        section_prompts = self._prompts.get("section_prompts", {})

        # 动态章节顺序：优先使用 reference_structure 中的章节（用户可增删）
        ref_sections_map = reference_structure.get("sections", {})
        if ref_sections_map:
            standard_keys = [k for k, _ in self._section_order]
            # 保持标准章节顺序，追加用户自定义 key
            ordered_keys = [k for k in standard_keys if k in ref_sections_map]
            ordered_keys += [k for k in ref_sections_map if k not in standard_keys]
            section_order = [
                (k, ref_sections_map[k].get("title", k) or dict(self._section_order).get(k, k))
                for k in ordered_keys
            ]
        else:
            section_order = self._section_order

        # 通用 prompt 模板（用于用户自定义章节）
        default_prompt = (
            "请为主题\"{theme}\"生成\"{section_title}\"章节内容。\n"
            "参考知识：{rag_context}\n"
            "网络资料：{web_context}\n"
            "参考内容：{reference_content}\n"
            "直接输出 Markdown，语言简洁专业。"
        )

        for section_key, section_title in section_order:
            await ws_callback({"type": "status", "content": f"正在生成章节：{section_title}"})

            # 获取该章节在参考文档中的内容
            ref_sections = reference_structure.get("sections", {})
            ref_section = ref_sections.get(section_key, {})
            reference_content = ref_section.get("content", "") if ref_section else ""

            # 构建 prompt（标准 key 用预设模板，自定义 key 用通用模板）
            prompt_template = section_prompts.get(section_key, default_prompt)
            # 使用 defaultdict 容错：模板中出现未知占位符时返回空字符串
            format_kwargs = defaultdict(str, {
                "theme": theme,
                "section_title": section_title,
                "rag_context": rag_context[:1500] if rag_context else "（无）",
                "web_context": web_context[:1000] if web_context else "（无）",
                "reference_content": reference_content[:800] if reference_content else "（无参考内容）",
                "modules_content": sections.get("modules", "")[:500],
            })
            prompt = prompt_template.format_map(format_kwargs)

            # 流式生成
            section_content = ""
            try:
                async for chunk in llm_factory.stream(
                    prompt=prompt,
                    system_prompt=system_prompt,
                    temperature=0.7,
                ):
                    section_content += chunk
                    await ws_callback({"type": "chunk", "content": chunk})
            except Exception as e:
                logger.warning(f"Section '{section_key}' generation error: {e}")
                # 保留空内容，继续下一章节
                await ws_callback({"type": "log", "content": f"章节 {section_title} 生成出错: {e}"})

            sections[section_key] = section_content
            await ws_callback({
                "type": "section",
                "section": section_key,
                "content": section_content,
            })

        # Step 4: 组合并导出（按实际生成的章节顺序）
        full_md = self._assemble_markdown(theme, sections, section_order)
        md_path = self._save_markdown(full_md)
        docx_path = self._export_docx(full_md)

        result = {
            "content": full_md,
            "md_path": str(md_path),
            "docx_path": str(docx_path),
        }

        # 课程大纲模式：额外导出 PDF
        if self.mode == "syllabus":
            pdf_path = self._convert_docx_to_pdf(docx_path)
            if pdf_path:
                result["pdf_path"] = str(pdf_path)

        return result

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    async def _rag_retrieve(self, theme: str, kb_name: str) -> str:
        """调用 RAGService 检索与主题相关的知识。"""
        try:
            from src.services.rag.service import RAGService

            service = RAGService()
            result = await service.search(
                query=theme,
                kb_name=kb_name,
                mode="hybrid",
            )
            return result.get("answer") or result.get("content") or ""
        except Exception as e:
            logger.warning(f"RAG retrieval failed for kb '{kb_name}': {e}")
            return ""

    async def _web_search(self, theme: str) -> str:
        """调用 web_search 服务（同步函数，在线程池中运行）。"""
        try:
            from src.services.search import web_search

            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(None, lambda: web_search(theme))
            return result.get("answer", "")
        except Exception as e:
            logger.warning(f"Web search failed: {e}")
            return ""

    def _assemble_markdown(
        self, theme: str, sections: dict[str, str],
        section_order: list[tuple[str, str]] | None = None,
    ) -> str:
        """将各章节内容组合为完整 Markdown 文档。"""
        order = section_order if section_order is not None else self._section_order
        suffix = "课程大纲" if self.mode == "syllabus" else "实习任务书"
        header = f"# {theme} — {suffix}\n\n"
        parts = [header]
        for section_key, _ in order:
            content = sections.get(section_key, "")
            if content.strip():
                parts.append(content.strip())
                parts.append("\n\n---\n\n")
        return "".join(parts)

    def _save_markdown(self, content: str) -> Path:
        md_path = self.output_dir / "generated_task.md"
        md_path.write_text(content, encoding="utf-8")
        return md_path

    def _export_docx(self, markdown_content: str) -> Path:
        """将 Markdown 内容导出为 .docx 文件（使用 python-docx）。"""
        docx_path = self.output_dir / "generated_task.docx"
        try:
            from docx import Document
            from docx.shared import Pt

            doc = Document()
            for line in markdown_content.splitlines():
                stripped = line.strip()
                if stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped == "---":
                    doc.add_paragraph("─" * 40)
                elif stripped:
                    doc.add_paragraph(stripped)
            doc.save(str(docx_path))
        except Exception as e:
            logger.warning(f"docx export failed: {e}. Saving empty placeholder.")
            docx_path.write_bytes(b"")

        return docx_path


    def _convert_docx_to_pdf(self, docx_path: Path) -> Path | None:
        """将 docx 转换为 PDF（优先 LibreOffice，回退 docx2pdf）。"""
        pdf_path = docx_path.with_suffix(".pdf")
        # 尝试 LibreOffice
        try:
            import subprocess
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf",
                 "--outdir", str(docx_path.parent), str(docx_path)],
                capture_output=True, text=True, timeout=60,
            )
            if pdf_path.exists():
                logger.info(f"PDF exported via LibreOffice: {pdf_path}")
                return pdf_path
        except Exception as e:
            logger.debug(f"LibreOffice conversion failed: {e}")

        # 尝试 docx2pdf
        try:
            from docx2pdf import convert
            convert(str(docx_path), str(pdf_path))
            if pdf_path.exists():
                logger.info(f"PDF exported via docx2pdf: {pdf_path}")
                return pdf_path
        except Exception as e:
            logger.debug(f"docx2pdf conversion failed: {e}")

        logger.warning("PDF conversion not available (install LibreOffice or docx2pdf)")
        return None


__all__ = [
    "TaskGenerator",
    "SECTION_ORDER_ZH", "SECTION_ORDER_EN",
    "SECTION_ORDER_SYLLABUS_ZH", "SECTION_ORDER_SYLLABUS_EN",
]
